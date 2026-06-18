from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "doc_assets"
STEP1_BASE = ASSETS / "step1_base.png"
STEP1_ANNOTATED = ASSETS / "step1_annotated.png"
STEP2 = ASSETS / "step2.png"
OUT = ROOT / "为冰箱锁授权方法.docx"


def add_arrow(draw, start, end, color, width=5, head_len=26, head_width=18):
    draw.line([start, end], fill=color, width=width)
    import math

    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    left = (
        end[0] - head_len * math.cos(angle) + head_width * math.sin(angle) / 2,
        end[1] - head_len * math.sin(angle) - head_width * math.cos(angle) / 2,
    )
    right = (
        end[0] - head_len * math.cos(angle) - head_width * math.sin(angle) / 2,
        end[1] - head_len * math.sin(angle) + head_width * math.cos(angle) / 2,
    )
    draw.polygon([end, left, right], fill=color)


def annotate_step1():
    img = Image.open(STEP1_BASE).convert("RGBA")
    draw = ImageDraw.Draw(img)
    red = (255, 77, 79, 255)

    # Recreate the navigation arrow and the card-number highlight visible in the supplied screenshot.
    add_arrow(draw, (278, 121), (181, 187), red, width=5)
    add_arrow(draw, (198, 198), (459, 466), red, width=5)
    draw.rectangle((525, 474, 1200, 533), outline=red, width=4)

    img.save(STEP1_ANNOTATED)


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def set_cell_borders(cell, color="D9E2EC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "8")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tc_pr.append(borders)


def style_paragraph(paragraph, font_name="Calibri", size=11, color="000000", bold=False):
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold


def add_step(doc, label, text, image_path):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E8EEF5")
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(label)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4D78")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4D78")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(5)

    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_after = Pt(0)
    pic_p.add_run().add_picture(str(image_path), width=Inches(6.5))


def build_doc():
    annotate_step1()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("为冰箱锁授权方法")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    add_step(
        doc,
        "1. ",
        "进入人员管理-人员资料，点击编辑，输入卡号/副卡号",
        STEP1_ANNOTATED,
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_step(
        doc,
        "2. ",
        "点击仪器管理-仪器授权，搜索需要授权的用户名称，再在右边搜索需授权的冰箱锁勾选即可",
        STEP2,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
